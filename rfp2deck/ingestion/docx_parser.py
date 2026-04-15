from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Union

from docx import Document


@dataclass
class ParsedDoc:
    text: str
    paragraph_count: int


def parse_docx(path_or_bytes: Union[Path, bytes]) -> ParsedDoc:
    if isinstance(path_or_bytes, bytes):
        d = Document(BytesIO(path_or_bytes))
    else:
        d = Document(path_or_bytes)
    paras = [p.text.strip() for p in d.paragraphs if p.text and p.text.strip()]
    return ParsedDoc(text="\n".join(paras), paragraph_count=len(paras))
